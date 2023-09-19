# -*- coding: utf-8 -*-
"""
文 件 名: copy_table.py
文件描述: 复制已有的表格（很困难，网上也搜索不到相关资料，无法实现）
https://blog.csdn.net/weixin_44312186/article/details/104944110
放弃还是在原来的基础上增加行吧
作    者: HanKin
创建日期: 2023.09.18
修改日期：2023.09.18
Copyright (c) 2023 HanKin. All rights reserved.
"""

import docx

doc = docx.Document('your_document.docx')
source_table = doc.tables[0]  # 假设要复制样式的源表格
target_table = doc.tables[1]  # 假设要应用样式的目标表格

# 复制表格样式
target_table.style = "Table Grid"

# 复制表格边框
for i, row in enumerate(source_table.rows):
    target_row = target_table.rows[i]
    for j, cell in enumerate(row.cells):
        target_cell = target_row.cells[j]
        target_cell._tc = cell._tc

# 成功
table = doc.add_table(5, 3, style="Table Grid")

# 失败
table = doc.add_table(5, 3, style=target_table.style)

doc.save('your_new_document.docx')
