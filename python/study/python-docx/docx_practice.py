# -*- coding: utf-8 -*-
"""
文 件 名: docx_practice.py
文件描述: 使用python-docx库实战，打开指定模板写排查报告
作    者: HanKin
创建日期: 2023.09.15
修改日期：2023.09.15
Copyright (c) 2023 HanKin. All rights reserved.
"""
import os
import time
from docx import Document       # 导入docx库
from docx.shared import Inches  # 导入英寸单位（可用于指定图片大小、表格宽高等）
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt
import json
import datetime
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

DOCX_FILE_PATH = "游戏助手排查报告模板.docx"
# 检测约定
detection_levels = {"info": "正常", "warning": "告警", "error": "异常"}
device_dir = "./"

def read_docx_table():
    """读取docx中所有表格内容，可以定位表格序号
    """

    document = Document(DOCX_FILE_PATH)

    # 读取表格内容
    for tb in document.tables:
        for i,row in enumerate(tb.rows):
            for j,cell in enumerate(row.cells):
                text = ""
                for p in cell.paragraphs:
                    text += p.text

                print(f"第{i}行，第{j}列的内容{text}")

def docx_table_add_text(cell, text, picture=None):
    """
    docx文档表格填写数据

    :return zip: 一个zip压缩包
    """

    #print(len(cell.paragraphs[0].runs))
    cell.text = ""
    run = cell.paragraphs[0].add_run()
    if picture:
        run.add_picture(os.path.join(device_dir, picture), height=Inches(0.12), width=Inches(0.12))

    run.add_text(text)
    font = run.font
    font.name = u"微软雅黑"
    font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), u"微软雅黑")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def generate_troubleshooting_report(TROUBLESHOOTING_REPORT):
    """
    生成游戏排查报告文档

    :return zip: 一个zip压缩包
    """

    # 打开报告模板
    document = Document(os.path.join(device_dir, DOCX_FILE_PATH))

    # 联系方式表格
    contact_table = document.tables[0]
    contact_table.cell(3, 1).text = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M-%S")

    # VDI信息表格
    vdi_table = document.tables[2]
    vdi_table.cell(0, 1).text = TROUBLESHOOTING_REPORT["game_version"]
    vdi_table.cell(1, 1).text = TROUBLESHOOTING_REPORT["phone_version"]
    vdi_table.cell(2, 1).text = TROUBLESHOOTING_REPORT["terminal_sn"]
    vdi_table.cell(3, 1).text = TROUBLESHOOTING_REPORT["terminal_version"]
    vdi_table.cell(4, 1).text = TROUBLESHOOTING_REPORT["user_name"]
    vdi_table.cell(5, 1).text = TROUBLESHOOTING_REPORT["server_version"]
    vdi_table.cell(6, 1).text = TROUBLESHOOTING_REPORT["server_version"]
    vdi_table.cell(7, 1).text = TROUBLESHOOTING_REPORT["server_type"]

    # 设备信息表格
    device_table = document.tables[3]
    device_table.cell(0, 1).text = TROUBLESHOOTING_REPORT["phone_name"]
    device_table.cell(1, 1).text = TROUBLESHOOTING_REPORT["phone_id"]
    device_table.cell(2, 1).text = TROUBLESHOOTING_REPORT["function1"]
    if "function2" in TROUBLESHOOTING_REPORT:
        device_table.cell(3, 1).text = TROUBLESHOOTING_REPORT["function2"]

    # 一键检测结果之检测项表格
    one_click_detection_table = document.tables[4]
    detection_list = TROUBLESHOOTING_REPORT["one_click_detection_list"]
    detection_num = len(detection_list)
    table_need_rows = detection_num + 1 # 增加行数
    for i in range(table_need_rows - len(one_click_detection_table.rows)):
        one_click_detection_table.add_row()
    warning_error_items = []
    for i in range(detection_num):
        docx_table_add_text(one_click_detection_table.cell(i+1, 0), str(i+1))
        docx_table_add_text(one_click_detection_table.cell(i+1, 1), detection_list[i]["content"])
        level = detection_list[i]["level"]
        docx_table_add_text(one_click_detection_table.cell(i+1, 2), " {}".format(detection_levels[level]), "{}.bmp".format(level))
        if level != "info":
            warning_error_items.append(detection_list[i])

    # 一键检测结果之异常信息表格
    warning_error_table = document.tables[5]
    warning_error_num = len(warning_error_items)
    table_need_rows = warning_error_num * 4 # 增加行数
    for i in range(table_need_rows - len(warning_error_table.rows)):
        warning_error_table.add_row()
    for i in range(warning_error_num):
        docx_table_add_text(warning_error_table.cell(i*4+0, 0), "检测项")
        docx_table_add_text(warning_error_table.cell(i*4+0, 1), warning_error_items[i]["content"])
        docx_table_add_text(warning_error_table.cell(i*4+1, 0), "结果")
        docx_table_add_text(warning_error_table.cell(i*4+1, 1),
                            " {}".format(detection_levels[warning_error_items[i]["level"]]),
                            "{}.bmp".format(warning_error_items[i]["level"]))
        docx_table_add_text(warning_error_table.cell(i*4+2, 0), "问题原因")
        docx_table_add_text(warning_error_table.cell(i*4+2, 1), warning_error_items[i]["cause"])
        docx_table_add_text(warning_error_table.cell(i*4+3, 0), "处置建议")
        docx_table_add_text(warning_error_table.cell(i*4+3, 1), warning_error_items[i]["suggestion"])
        warning_error_table.cell(i*4, 1)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'),bgColor = RGBColor(153, 204, 255))))
    
    # 人工排查表格
    manual_troubleshooting_table = document.tables[6]
    troubleshooting_scheme = TROUBLESHOOTING_REPORT["troubleshooting_scheme"]
    troubleshooting_scheme_num = len(troubleshooting_scheme)
    table_need_rows = 0
    for i in range(troubleshooting_scheme_num):
        troubleshooting_steps = len(troubleshooting_scheme[i]["troubleshooting"])
        start_row = table_need_rows # 每个问题场景起始的行序号
        table_need_rows += troubleshooting_steps+1 # 增加行数
        for j in range(table_need_rows - len(manual_troubleshooting_table.rows)):
            manual_troubleshooting_table.add_row()
        manual_troubleshooting_table.cell(start_row, 0).merge(manual_troubleshooting_table.cell(start_row, 1))
        docx_table_add_text(manual_troubleshooting_table.cell(start_row, 0), "问题场景")
        docx_table_add_text(manual_troubleshooting_table.cell(start_row, 2), troubleshooting_scheme[i]["scenario"])
        manual_troubleshooting_table.cell(start_row, 2)._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="{bgColor}"/>'.format(nsdecls('w'),bgColor = RGBColor(153, 204, 255))))
        troubleshooting = troubleshooting_scheme[i]["troubleshooting"]
        for j in range(troubleshooting_steps):
            docx_table_add_text(manual_troubleshooting_table.cell(start_row+j+1, 0), troubleshooting[j]["step"])
            docx_table_add_text(manual_troubleshooting_table.cell(start_row+j+1, 1), troubleshooting[j]["item"])
            docx_table_add_text(manual_troubleshooting_table.cell(start_row+j+1, 2), troubleshooting[j]["operation"])

    # 保存排查报告
    timestamp = datetime.datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
    document.save(os.path.join(device_dir, "{}_游戏助手排查报告.docx".format(timestamp)))

def main():
    """主函数
    """

    read_docx_table()

def debug():
    """调试函数
    """
    
    TROUBLESHOOTING_REPORT = dict()
    TROUBLESHOOTING_REPORT["game_version"] = "2018.08.05"
    TROUBLESHOOTING_REPORT["phone_version"] = "mi10pro"
    TROUBLESHOOTING_REPORT["terminal_version"] = "2023.09.18"
    TROUBLESHOOTING_REPORT["terminal_sn"] = "98371982bnbmnvsdf"
    TROUBLESHOOTING_REPORT["user_name"] = "异次元的狙击手"
    TROUBLESHOOTING_REPORT["server_version"] = "一区"
    TROUBLESHOOTING_REPORT["server_type"] = "开服区"
    TROUBLESHOOTING_REPORT["phone_name"] = "小米"
    TROUBLESHOOTING_REPORT["phone_id"] = "10pro"
    TROUBLESHOOTING_REPORT["function1"] = "打游戏"
    
    one_click_detection_list = []
    info_result = {"level": "info", "content": "游戏登录是否存在问题", "cause": "正常登录",
                    "suggestion": ""}
    warning_result = {"level": "warning", "content": "苍牙是否能释放大招", "cause": "有问题",
                    "suggestion": "不影响使用"}
    error_result = {"level": "error", "content": "钥匙抽奖机制怎么样", "cause": "长期抽不到理想的SSR",
                    "suggestion": "更改机制"}
    one_click_detection_list.append(info_result)
    one_click_detection_list.append(warning_result)
    one_click_detection_list.append(error_result)
    TROUBLESHOOTING_REPORT["one_click_detection_list"] = one_click_detection_list
    
    troubleshooting_scheme = []
    try:
        with open(os.path.join(device_dir, "manualTroubleshooting.json"), encoding="utf-8") as f:
            manual_troubleshooting_dic = json.load(f)
    except:
        print("open manualTroubleshooting.json file failed")
        print(traceback.format_exc())
        return troubleshooting_scheme
    for ninja in ["苍牙", "卫鲤"]:
        troubleshooting_scheme.extend(manual_troubleshooting_dic.get(ninja))
    TROUBLESHOOTING_REPORT["troubleshooting_scheme"] = troubleshooting_scheme

    generate_troubleshooting_report(TROUBLESHOOTING_REPORT)

if __name__ == '__main__':
    """程序入口
    """
    
    #os.system('chcp 936 & cls')
    print('******** starting ********')
    start_time = time.time()
    
    #main()
    debug()

    end_time = time.time()
    print('process spend {} s.\n'.format(round(end_time - start_time, 3)))