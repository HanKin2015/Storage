# -*- coding: utf-8 -*-
"""
文 件 名: docx_example.py
文件描述: 学习python-docx库
备    注：https://mp.weixin.qq.com/s/PoLoFs0M6I3ftu_uRiUmsw
作    者: HanKin
创建日期: 2023.09.15
修改日期：2023.09.15
Copyright (c) 2023 HanKin. All rights reserved.
"""
import os
import time
from docx import Document       # 导入docx库
from docx.shared import Inches  # 导入英寸单位（可用于指定图片大小、表格宽高等）
from docx.enum.style import WD_STYLE_TYPE

def test():
    """测试函数
    """
    
    # 新建一个空文档
    document = Document()
    document.save('new.docx')
    
    # 操作旧文档
    document = Document('exist.docx')
    document.save('new.docx')

def tutor():
    """Python-docx官方例程解析
    """

    # 新建一个空文档
    document = Document()

    # 设置标题段落
    document.add_heading('Document Title', 0)

    # 添加段落
    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True       # 在指定段落后添加粗体文字
    p.add_run(' and some ')             # 在指定段落后添加默认格式文字
    p.add_run('italic.').italic = True  # 在指定段落后添加斜体文字

    # 添加1级标题=标题1
    document.add_heading('Heading, level 1', level=1)

    # 添加指定格式段落
    document.add_paragraph('Intense quote', style='Intense Quote')
    document.add_paragraph('first item in unordered list', style='List Bullet')
    document.add_paragraph('first item in ordered list', style='List Number')

    # 添加图片
    document.add_picture('t01d1d535f53114cbb7.jpg', width=Inches(1.25))

    # 待添加如表格的内容
    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    # 添加一个1行3列的表格
    table = document.add_table(rows=1, cols=3)
    
    # 填充标题行
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    
    # 填充数据行
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    # 添加分页符
    document.add_page_break()

    # 保存文档
    document.save('4.1 Python-docx官方例程.docx')

def table_style():
    """Python-docx 表格样式设置
    """
    
    document = Document()

    table = document.add_table(3, 3, style="Medium Grid 1 Accent 1")
    heading_cells = table.rows[0].cells
    heading_cells[0].text = '第一列内容'
    heading_cells[1].text = '第二列内容'
    heading_cells[2].text = '第三列内容'

    document.save("demo.docx")


def table_all_styles():
    """遍历所有样式
    """
    
    document = Document()
    styles = document.styles

    # 生成所有表样式
    for s in styles:
        if s.type == WD_STYLE_TYPE.TABLE:
            document.add_paragraph("表格样式 :  " + s.name)
            table = document.add_table(3, 3, style=s)
            heading_cells = table.rows[0].cells
            heading_cells[0].text = '第一列内容'
            heading_cells[1].text = '第二列内容'
            heading_cells[2].text = '第三列内容'
            document.add_paragraph("\n")

    document.save('4.3 所有表格样式.docx')

def operate_old_docx():
    """Python-docx 修改旧 word 文档
    读取word文档的内容
    """
    
    document = Document('外设助手排查报告模板.docx')

    # 读取 word 中所有内容
    for p in document.paragraphs:
        print("paragraphs：", p.text)

    # 读取 word 中所有一级标题
    for p in document.paragraphs:
        if p.style.name == 'Heading 1':
            print("Heading 1：", p.text)

    # 读取 word 中所有二级标题
    for p in document.paragraphs:
        if p.style.name == 'Heading 2':
            print("Heading 2：", p.text)

    # 读取 word 中所有正文
    for p in document.paragraphs:
        if p.style.name == 'Normal':
            print("Normal：", p.text)

    document.save('修改后的报告.docx')

def read_docx_table():
    """读取docx中表格内容
    """

    document = Document('外设助手排查报告模板.docx')

    # 读取表格内容
    for tb in document.tables:
        for i,row in enumerate(tb.rows):
            for j,cell in enumerate(row.cells):
                text = ''
                for p in cell.paragraphs:
                    text += p.text

                print(f'第{i}行，第{j}列的内容{text}')

def main():
    """主函数
    """
    
    #test()
    #table_style()
    #table_all_styles()
    #operate_old_docx()
    read_docx_table()

if __name__ == '__main__':
    """程序入口
    """
    
    #os.system('chcp 936 & cls')
    print('******** starting ********')
    start_time = time.time()
    
    main()

    end_time = time.time()
    print('process spend {} s.\n'.format(round(end_time - start_time, 3)))
